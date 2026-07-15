from __future__ import annotations

import logging
import math
from pathlib import Path
from tempfile import NamedTemporaryFile
from urllib.parse import urlparse
from uuid import NAMESPACE_URL, uuid5

import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

LOGGER = logging.getLogger(__name__)


class DocumentProcessor:
    supported_extensions = {".docx", ".txt", ".md", ".rst", ".pdf", ".xls", ".xlsx"}

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50) -> None:
        self.chunk_size = chunk_size
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""],
            length_function=len,
        )

    def load(self, source: str) -> list[Document]:
        source = source.strip().strip('"')
        if source.startswith(("http://", "https://")):
            documents = self._load_url(source)
        else:
            path = Path(source).expanduser().resolve()
            if path.is_file():
                documents = self._load_file(path)
            elif path.is_dir():
                documents = self._load_directory(path)
            else:
                raise FileNotFoundError(f"Файл или папка не найдены: {source}")

        chunks: list[Document] = []
        for document in documents:
            metadata = dict(document.metadata)
            if metadata.pop("_prechunked", False):
                chunks.append(Document(page_content=document.page_content, metadata=metadata))
            else:
                chunks.extend(
                    self.splitter.split_documents(
                        [Document(page_content=document.page_content, metadata=metadata)]
                    )
                )
        unique_chunks: list[Document] = []
        seen_chunks: set[tuple[str, str, str]] = set()
        for chunk in chunks:
            key = (
                str(chunk.metadata.get("source", "unknown")),
                str(chunk.metadata.get("page", chunk.metadata.get("sheet", ""))),
                chunk.page_content,
            )
            if key not in seen_chunks:
                seen_chunks.add(key)
                unique_chunks.append(chunk)

        for index, chunk in enumerate(unique_chunks):
            identity = "|".join(
                [
                    str(chunk.metadata.get("source", "unknown")),
                    str(chunk.metadata.get("page", chunk.metadata.get("sheet", ""))),
                    str(index),
                    chunk.page_content,
                ]
            )
            chunk.metadata["doc_id"] = str(uuid5(NAMESPACE_URL, identity))
            chunk.metadata["chunk_index"] = index
        return unique_chunks

    def _load_directory(self, directory: Path) -> list[Document]:
        documents: list[Document] = []
        for path in sorted(directory.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in self.supported_extensions:
                continue
            LOGGER.info("Обработка %s", path)
            try:
                documents.extend(self._load_file(path))
            except Exception as error:
                LOGGER.warning("Не удалось обработать %s: %s", path, error)
        return documents

    def _load_file(self, path: Path, source: str | None = None) -> list[Document]:
        suffix = path.suffix.lower()
        display_source = source or str(path)
        if suffix in {".txt", ".md", ".rst"}:
            return [
                Document(page_content=self._read_text(path), metadata={"source": display_source})
            ]
        if suffix == ".pdf":
            return self._load_pdf(path, display_source)
        if suffix == ".docx":
            return self._load_docx(path, display_source)
        if suffix in {".xls", ".xlsx"}:
            return self._load_excel(path, display_source)
        raise ValueError(f"Формат {suffix or '<без расширения>'} не поддерживается")

    @staticmethod
    def _read_text(path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp1251"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _load_pdf(path: Path, source: str) -> list[Document]:
        reader = PdfReader(str(path))
        documents = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": source, "page": page_number},
                    )
                )
        return documents

    @staticmethod
    def _load_docx(path: Path, source: str) -> list[Document]:
        document = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        return [Document(page_content=text, metadata={"source": source})]

    def _load_excel(self, path: Path, source: str) -> list[Document]:
        workbook = pd.ExcelFile(path)
        documents: list[Document] = []
        for sheet_name in workbook.sheet_names:
            raw_frame = pd.read_excel(workbook, sheet_name=sheet_name, header=None, dtype=object)
            if raw_frame.empty:
                continue
            raw_frame = self._normalize_excel_frame(raw_frame)
            for block_index, bounds in enumerate(self._excel_blocks(raw_frame), start=1):
                documents.extend(
                    self._excel_block_documents(
                        raw_frame,
                        bounds,
                        source,
                        sheet_name,
                        block_index,
                    )
                )
        return documents

    @classmethod
    def _normalize_excel_frame(cls, frame: pd.DataFrame) -> pd.DataFrame:
        """Collapse repeated values used to imitate merged cells in generated workbooks."""
        normalized = frame.copy()
        previous_single_text: tuple[int, str] | None = None

        for row_position in range(len(normalized)):
            row = normalized.iloc[row_position]
            populated = [
                (column_position, value)
                for column_position, value in enumerate(row.tolist())
                if cls._cell_has_value(value)
            ]
            repeated_text = (
                len(populated) > 1
                and all(isinstance(value, str) for _, value in populated)
                and len({str(value).strip() for _, value in populated}) == 1
            )
            if repeated_text:
                for column_position, _ in populated[1:]:
                    normalized.iat[row_position, column_position] = pd.NA
                populated = populated[:1]

            if len(populated) == 1 and isinstance(populated[0][1], str):
                signature = (populated[0][0], str(populated[0][1]).strip())
                if signature == previous_single_text:
                    normalized.iloc[row_position, :] = pd.NA
                    continue
                previous_single_text = signature
            else:
                previous_single_text = None
        return normalized

    @classmethod
    def _excel_blocks(cls, frame: pd.DataFrame) -> list[tuple[int, int, int, int]]:
        row_flags = [
            any(cls._cell_has_value(value) for value in row.tolist()) for _, row in frame.iterrows()
        ]
        blocks: list[tuple[int, int, int, int]] = []
        for first_row, last_row in cls._contiguous_ranges(row_flags):
            row_band = frame.iloc[first_row : last_row + 1]
            column_flags = [
                any(cls._cell_has_value(value) for value in row_band.iloc[:, column].tolist())
                for column in range(frame.shape[1])
            ]
            for first_column, last_column in cls._contiguous_ranges(column_flags):
                blocks.append((first_row, last_row, first_column, last_column))
        return blocks

    def _excel_block_documents(
        self,
        frame: pd.DataFrame,
        bounds: tuple[int, int, int, int],
        source: str,
        sheet_name: str,
        block_index: int,
    ) -> list[Document]:
        first_row, last_row, first_column, last_column = bounds
        block = frame.iloc[first_row : last_row + 1, first_column : last_column + 1]
        header_offset = self._detect_excel_header(block)
        documents: list[Document] = []

        if header_offset is not None and header_offset > 0:
            prefix = block.iloc[:header_offset]
            documents.extend(
                self._excel_freeform_documents(
                    prefix,
                    source,
                    sheet_name,
                    block_index,
                    first_column,
                    last_column,
                )
            )

        if header_offset is None:
            return self._excel_freeform_documents(
                block,
                source,
                sheet_name,
                block_index,
                first_column,
                last_column,
            )

        header_row = block.iloc[header_offset]
        headers = self._unique_headers(header_row.tolist())
        data_frame = block.iloc[header_offset + 1 :].copy()
        data_frame.columns = headers
        data_frame = data_frame.dropna(axis="rows", how="all")
        if data_frame.empty:
            return documents

        if self._is_key_value_headers(headers):
            entries = [
                (
                    f"{self._format_cell(row.iloc[0])}: {self._format_cell(row.iloc[1])}",
                    int(row_index) + 1,
                )
                for row_index, row in data_frame.iterrows()
                if self._cell_has_value(row.iloc[0]) and self._cell_has_value(row.iloc[1])
            ]
            documents.extend(
                self._excel_lines_documents(
                    source,
                    sheet_name,
                    entries,
                    block_index=block_index,
                    block_type="key_value",
                    first_column=first_column,
                    last_column=last_column,
                    heading="Справочные факты (поле: значение):",
                )
            )
            return documents

        documents.extend(
            self._excel_table_documents(
                data_frame,
                headers,
                source,
                sheet_name,
                block_index,
                first_column,
                last_column,
            )
        )
        documents.extend(
            self._excel_aggregate_documents(
                data_frame,
                source,
                sheet_name,
                block_index,
                first_column,
                last_column,
            )
        )
        return documents

    def _excel_table_documents(
        self,
        data_frame: pd.DataFrame,
        headers: list[str],
        source: str,
        sheet_name: str,
        block_index: int,
        first_column: int,
        last_column: int,
    ) -> list[Document]:
        entries: list[tuple[str, int]] = []
        for row_index, row in data_frame.iterrows():
            fields = [
                f"{column}: {self._format_cell(value)}"
                for column, value in row.items()
                if self._cell_has_value(value)
            ]
            if fields:
                row_number = int(row_index) + 1
                entries.append((f"Строка Excel {row_number}: " + "; ".join(fields), row_number))
        return self._excel_lines_documents(
            source,
            sheet_name,
            entries,
            block_index=block_index,
            block_type="table",
            first_column=first_column,
            last_column=last_column,
            heading="Колонки: " + " | ".join(headers),
        )

    def _excel_freeform_documents(
        self,
        block: pd.DataFrame,
        source: str,
        sheet_name: str,
        block_index: int,
        first_column: int,
        last_column: int,
    ) -> list[Document]:
        is_key_value = block.shape[1] == 2 and all(
            sum(self._cell_has_value(value) for value in row.tolist()) == 2
            for _, row in block.iterrows()
        )
        entries: list[tuple[str, int]] = []
        for row_index, row in block.iterrows():
            values = [value for value in row.tolist() if self._cell_has_value(value)]
            if not values:
                continue
            if is_key_value:
                line = f"{self._format_cell(values[0])}: {self._format_cell(values[1])}"
            else:
                line = " | ".join(self._format_cell(value) for value in values)
            entries.append((line, int(row_index) + 1))
        return self._excel_lines_documents(
            source,
            sheet_name,
            entries,
            block_index=block_index,
            block_type="key_value" if is_key_value else "text",
            first_column=first_column,
            last_column=last_column,
            heading="Сводные показатели (показатель: значение):" if is_key_value else "Текст:",
        )

    def _excel_lines_documents(
        self,
        source: str,
        sheet_name: str,
        entries: list[tuple[str, int]],
        *,
        block_index: int,
        block_type: str,
        first_column: int,
        last_column: int,
        heading: str,
    ) -> list[Document]:
        if not entries:
            return []
        prefix = [f'Лист Excel: "{sheet_name}"', heading]
        documents: list[Document] = []
        current_lines = list(prefix)
        first_row: int | None = None
        last_row: int | None = None

        for line, row_number in entries:
            candidate = "\n".join([*current_lines, line])
            if len(candidate) > self.chunk_size and len(current_lines) > len(prefix):
                documents.append(
                    self._excel_document(
                        source,
                        sheet_name,
                        current_lines,
                        first_row,
                        last_row,
                        block_index=block_index,
                        block_type=block_type,
                        first_column=first_column,
                        last_column=last_column,
                    )
                )
                current_lines = list(prefix)
                first_row = None
            current_lines.append(line)
            first_row = first_row or row_number
            last_row = row_number

        documents.append(
            self._excel_document(
                source,
                sheet_name,
                current_lines,
                first_row,
                last_row,
                block_index=block_index,
                block_type=block_type,
                first_column=first_column,
                last_column=last_column,
            )
        )
        return documents

    def _excel_aggregate_documents(
        self,
        data_frame: pd.DataFrame,
        source: str,
        sheet_name: str,
        block_index: int,
        first_column: int,
        last_column: int,
    ) -> list[Document]:
        documents: list[Document] = []
        for column in data_frame.columns:
            source_series = data_frame[column]
            non_empty_count = sum(self._cell_has_value(value) for value in source_series)
            numeric = pd.to_numeric(source_series, errors="coerce").dropna()
            if len(numeric) < 2 or len(numeric) < math.ceil(non_empty_count * 0.8):
                continue

            minimum_row = int(numeric.idxmin())
            maximum_row = int(numeric.idxmax())
            lines = [
                f'Лист Excel: "{sheet_name}"',
                f'Агрегаты по всем строкам столбца "{column}":',
                self._aggregate_scope(data_frame, str(column)),
                f"Количество числовых значений: {len(numeric)}",
                (
                    f'Минимум по показателю "{column}": '
                    f"{self._format_number(numeric.loc[minimum_row])} "
                    f"(строка Excel {minimum_row + 1}{self._row_description(data_frame.loc[minimum_row], column)})"
                ),
                (
                    f'Максимум по показателю "{column}": '
                    f"{self._format_number(numeric.loc[maximum_row])} "
                    f"(строка Excel {maximum_row + 1}{self._row_description(data_frame.loc[maximum_row], column)})"
                ),
                f'Сумма по показателю "{column}": {self._format_number(numeric.sum())}',
                f'Среднее по показателю "{column}": {self._format_number(numeric.mean())}',
            ]
            documents.append(
                self._excel_document(
                    source,
                    sheet_name,
                    lines,
                    int(data_frame.index.min()) + 1,
                    int(data_frame.index.max()) + 1,
                    block_index=block_index,
                    block_type="aggregate",
                    first_column=first_column,
                    last_column=last_column,
                    extra_metadata={"aggregate_column": str(column)},
                )
            )
        return documents

    @staticmethod
    def _aggregate_scope(data_frame: pd.DataFrame, column: str) -> str:
        first_column = str(data_frame.columns[0]).strip().lower()
        if "месяц" in first_column:
            return (
                f'Уровень агрегации: месяц; максимальная месячная величина показателя "{column}" '
                "сравнивается по итогам месяцев."
            )
        if first_column in {"id", "ид", "номер"} or first_column.endswith(" id"):
            return (
                "Уровень агрегации: отдельная запись; максимум относится к одной строке, "
                "а не к месячному или общему итогу."
            )
        if any(token in first_column for token in ("время", "дата", "час")):
            return "Уровень агрегации: отдельный момент времени или измерение."
        return "Уровень агрегации: строки текущей таблицы."

    @classmethod
    def _detect_excel_header(cls, frame: pd.DataFrame) -> int | None:
        header_hints = {
            "поле",
            "описание",
            "показатель",
            "значение",
            "параметр",
            "id",
            "дата",
            "месяц",
            "клиент",
            "статус",
        }
        for offset, (_, row) in enumerate(frame.head(10).iterrows()):
            values = [value for value in row.tolist() if cls._cell_has_value(value)]
            if len(values) < 2 or not all(isinstance(value, str) for value in values):
                continue
            following = frame.iloc[offset + 1 :]
            if following.empty or not any(
                sum(cls._cell_has_value(value) for value in item.tolist()) >= 2
                for _, item in following.iterrows()
            ):
                continue
            normalized = {str(value).strip().lower() for value in values}
            following_has_non_text = any(
                cls._cell_has_value(value) and not isinstance(value, str)
                for value in following.to_numpy().ravel().tolist()
            )
            if len(values) >= 3 or normalized & header_hints or following_has_non_text:
                return offset
        return None

    @staticmethod
    def _is_key_value_headers(headers: list[str]) -> bool:
        if len(headers) != 2:
            return False
        first = headers[0].strip().lower()
        second = headers[1].strip().lower()
        return first in {"поле", "показатель", "параметр", "метрика"} and second in {
            "описание",
            "значение",
            "value",
        }

    @staticmethod
    def _contiguous_ranges(flags: list[bool]) -> list[tuple[int, int]]:
        ranges: list[tuple[int, int]] = []
        start: int | None = None
        for index, enabled in enumerate([*flags, False]):
            if enabled and start is None:
                start = index
            elif not enabled and start is not None:
                ranges.append((start, index - 1))
                start = None
        return ranges

    @staticmethod
    def _cell_has_value(value: object) -> bool:
        return not pd.isna(value) and bool(str(value).strip())

    @staticmethod
    def _unique_headers(values: list[object]) -> list[str]:
        headers: list[str] = []
        counts: dict[str, int] = {}
        for index, value in enumerate(values, start=1):
            base = str(value).strip() if not pd.isna(value) else f"column_{index}"
            count = counts.get(base, 0) + 1
            counts[base] = count
            headers.append(base if count == 1 else f"{base}_{count}")
        return headers

    @staticmethod
    def _format_cell(value: object) -> str:
        if isinstance(value, pd.Timestamp):
            return value.isoformat(sep=" ")
        if isinstance(value, float):
            return f"{value:.15g}"
        return str(value)

    @staticmethod
    def _format_number(value: object) -> str:
        number = float(value)
        if number.is_integer():
            return f"{int(number):,}".replace(",", " ")
        formatted = f"{number:,.6f}".rstrip("0").rstrip(".")
        if "." not in formatted:
            return formatted.replace(",", " ")
        integer, decimal = formatted.split(".")
        return f"{integer.replace(',', ' ')}.{decimal}"

    @classmethod
    def _row_description(cls, row: pd.Series, aggregate_column: str) -> str:
        fields: list[str] = []
        for column, value in row.items():
            if column == aggregate_column or not cls._cell_has_value(value):
                continue
            if isinstance(value, (int, float)) and not isinstance(value, bool):
                continue
            fields.append(f"{column}: {cls._format_cell(value)}")
            if len(fields) == 6:
                break
        return "; " + "; ".join(fields) if fields else ""

    @staticmethod
    def _excel_document(
        source: str,
        sheet_name: str,
        lines: list[str],
        first_row: int | None,
        last_row: int | None,
        *,
        block_index: int,
        block_type: str,
        first_column: int,
        last_column: int,
        extra_metadata: dict[str, object] | None = None,
    ) -> Document:
        metadata: dict[str, object] = {
            "source": source,
            "sheet": sheet_name,
            "block_index": block_index,
            "block_type": block_type,
            "first_row": first_row,
            "last_row": last_row,
            "first_column": DocumentProcessor._excel_column_name(first_column),
            "last_column": DocumentProcessor._excel_column_name(last_column),
            "_prechunked": True,
        }
        if extra_metadata:
            metadata.update(extra_metadata)
        return Document(
            page_content="\n".join(lines),
            metadata=metadata,
        )

    @staticmethod
    def _excel_column_name(index: int) -> str:
        name = ""
        number = index + 1
        while number:
            number, remainder = divmod(number - 1, 26)
            name = chr(65 + remainder) + name
        return name

    def _load_url(self, url: str) -> list[Document]:
        response = requests.get(url, timeout=60)
        response.raise_for_status()
        suffix = Path(urlparse(url).path).suffix.lower()
        if suffix in self.supported_extensions:
            with NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                temp_file.write(response.content)
                temp_path = Path(temp_file.name)
            try:
                return self._load_file(temp_path, source=url)
            finally:
                temp_path.unlink(missing_ok=True)

        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
        return [Document(page_content=text, metadata={"source": url})]
